from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH


def create_git_conflict_doc():
    doc = Document()

    # 1. 标题
    title = doc.add_heading('Git 冲突解决实战指南', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # 2. 为什么会产生冲突？
    doc.add_heading('一、 冲突的本质', level=1)
    doc.add_paragraph(
        '冲突不是错误，而是 Git 的一种“保护机制”。当两个分支修改了【同一个文件的同一行】，'
        '或者一个分支删除了文件而另一个分支修改了它时，Git 无法自动判断该保留谁的代码，于是停下来请你人工裁决。'
    )

    # 3. 核心解决流程
    doc.add_heading('二、 标准解决五步法', level=1)

    steps = [
        ("1. 触发冲突", "当你执行 git pull 或 git merge 时，终端出现 'CONFLICT' 字样。"),
        ("2. 定位文件", "输入 git status，状态为 'both modified' 的文件就是冲突文件。"),
        ("3. 打开编辑器",
         "在 IDEA 或 VS Code 中打开文件，寻找特殊标记：\n<<<<<<< HEAD (当前分支内容)\n=======\n>>>>>>> (要合并的内容)"),
        ("4. 人工裁决", "删除这些标记符号，并将代码修改为你最终想要的样子。"),
        ("5. 标记解决", "执行 git add <文件名> 告诉 Git 冲突已处理。")
    ]

    for step_title, step_desc in steps:
        p = doc.add_paragraph(style='List Number')
        run = p.add_run(step_title)
        run.bold = True
        doc.add_paragraph(step_desc)

    # 4. 完成合并
    doc.add_heading('三、 提交解决结果', level=1)
    doc.add_paragraph('完成所有文件的修改和 add 后，执行提交：')
    doc.add_paragraph('git commit -m "fix: resolve merge conflicts"', style='Quote')

    # 5. 预防冲突的“金科玉律”
    doc.add_heading('四、 减少冲突的职场习惯', level=1)
    tips = [
        "早点拉取：养成每天写代码前先 git pull 的习惯。",
        "小步提交：不要积攒了一周的代码才合并，提交越频繁，冲突范围越小。",
        "分工明确：尽量避免两个人在同一时间修改同一个文件的同一部分。"
    ]
    for tip in tips:
        doc.add_paragraph(tip, style='List Bullet')

    # 6. 后悔药：撤销合并
    doc.add_heading('五、 搞砸了怎么办？', level=1)
    doc.add_paragraph('如果解决过程中改乱了，想退回合并前的状态：')
    doc.add_paragraph('git merge --abort', style='Quote')

    # 保存
    file_name = 'Git冲突解决实战指南.docx'
    doc.save(file_name)
    print(f'成功！冲突指南已生成：{file_name}')


if __name__ == "__main__":
    create_git_conflict_doc()
